from fastapi import FastAPI, Path, HTTPException, Form
from fastapi.responses import FileResponse
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx2pdf import convert
from pathlib import Path
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
import os
import re


app = FastAPI()

# Allow cross-origin requests (CORS) for local development
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost", "http://localhost:5000", "izindteti.daniyalkautsar.com"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


app.mount("/static", StaticFiles(directory="static"), name="static")


def replace_text(doc, old_text, new_text, font_size=11):
    """
    Replace specified text in a Word document with new text while preserving formatting.

    Args:
        doc (Document): The Word document object.
        old_text (str): The text to be replaced.
        new_text (str): The new text to replace with.
        font_size (int, optional): The font size of the new text. Defaults to 11.
    """
    for paragraph in doc.paragraphs:
        if old_text in paragraph.text:
            for run in paragraph.runs:
                if old_text in run.text:
                    # Preserve formatting
                    font = run.font
                    run.clear()  # Clear existing run content
                    run.text = str(new_text)
                    run.bold = font.bold
                    run.italic = font.italic
                    run.underline = font.underline
                    run.font.size = Pt(font_size)
                    run.font.color.rgb = font.color.rgb

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if old_text in run.text:
                            # Preserve formatting
                            font = run.font
                            run.clear()  # Clear existing run content
                            run.text = new_text
                            run.bold = font.bold
                            run.italic = font.italic
                            run.underline = font.underline
                            run.font.size = Pt(font_size)
                            run.font.color.rgb = font.color.rgb


def save_as_word(doc, filename):
    """
    Save a Word document to a specified file.

    Args:
        doc (Document): The Word document object.
        filename (str): The filename to save the document as.
    """
    doc.save(filename)


def save_as_pdf(docx_filename, pdf_filename):
    """
    Convert a Word document to PDF format.

    Args:
        docx_filename (str): The filename of the Word document.
        pdf_filename (str): The filename to save the PDF document as.
    """
    convert(docx_filename, pdf_filename)


@app.get("/")
async def read_index():
    """
    Read the index page and return the static HTML file.
    """
    return FileResponse("static/index.html")


@app.post("/replace-and-download")
async def replace_and_download(
    namaMataKuliah: str = Form(...),
    kelasMataKuliah: str = Form(...),
    hariMataKuliah: str = Form(...),
    jamMataKuliah: str = Form(...),
    ruangMataKuliah: str = Form(...),
    namaDosen: str = Form(...),
    namaLengkap: str = Form(...),
    prodi: str = Form(...),
    nim: str = Form(...),
    nomorHP: str = Form(...),
    waktuIzin: str = Form(...),
    alasan: str = Form(...),
    waktuPermohonan: str = Form(...)
):
    """
    Replace text placeholders in a template Word document and download the modified document.

    Args:
        namaMataKuliah (str): The name of the course.
        kelasMataKuliah (str): The class of the course.
        hariMataKuliah (str): The day of the course.
        jamMataKuliah (str): The time of the course.
        ruangMataKuliah (str): The room of the course.
        namaDosen (str): The name of the lecturer.
        namaLengkap (str): The full name of the student.
        prodi (str): The study program of the student.
        nim (str): The student's ID number.
        nomorHP (str): The student's phone number.
        waktuIzin (str): The time of the permission.
        alasan (str): The reason for permission.
        waktuPermohonan (str): The time of the request.
    """
    # Assume the files are stored in a directory named "files"
    file_path = Path(f"files/template_surat_izin.docx")

    # Check if the file exists
    if not file_path.is_file():
        raise HTTPException(status_code=404, detail="File not found")

    # Load the Word document
    doc = Document(file_path)

    # Perform additional replacements based on the new input fields
    replace_text(doc, "nama_mata_kuliah", namaMataKuliah)
    replace_text(doc, "kelas_mata_kuliah", kelasMataKuliah)
    replace_text(doc, "jam_mata_kuliah", jamMataKuliah)
    replace_text(doc, "ruang_mata_kuliah", ruangMataKuliah)
    replace_text(doc, "nama_dosen", namaDosen)
    replace_text(doc, "nmhs", namaLengkap)
    replace_text(doc, "nomor_hp", nomorHP)
    replace_text(doc, "waktu_izin", waktuIzin)
    replace_text(doc, "waktu_permohonan", waktuPermohonan)

    for i in range(2):
        replace_text(doc, "day", hariMataKuliah)
        replace_text(doc, "prdi", prodi)
        replace_text(doc, "nims", nim)
        replace_text(doc, "ala_san", alasan)

    # Save the modified document as a Word file
    file_name = f"tmp/{namaLengkap}-{namaMataKuliah}-{waktuIzin}.docx"
    save_as_word(doc, file_name)

    return FileResponse(file_name, filename=file_name)
